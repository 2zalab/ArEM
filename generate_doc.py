from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)

NAVY  = RGBColor(0x00, 0x40, 0xA0)
BLUE  = RGBColor(0x5A, 0xC8, 0xFA)
BLACK = RGBColor(0x00, 0x00, 0x00)
GRAY  = RGBColor(0x55, 0x55, 0x55)

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_font(run, size, bold=False, italic=False, color=BLACK):
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    run.font.color.rgb = color

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    set_font(run, 16, bold=True, color=NAVY)
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '0040A0')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, 13, bold=True, color=NAVY)
    return p

def heading3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    set_font(run, 12, bold=True, color=GRAY)
    return p

def body(text, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    set_font(run, 11)
    return p

def body_bold(label, text, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    r1 = p.add_run(label)
    set_font(r1, 11, bold=True)
    r2 = p.add_run(text)
    set_font(r2, 11)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    run = p.add_run(text)
    set_font(run, 11)
    return p

def caption(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, 10, italic=True, color=GRAY)
    return p

def placeholder_image(label):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'[Capture d\'écran : {label}]')
    set_font(run, 10, italic=True, color=RGBColor(0xAA, 0xAA, 0xAA))
    return p

# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('CHAPITRE IV')
set_font(r, 14, bold=True, color=GRAY)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8)
r = p.add_run('Implémentation et Résultats Obtenus')
set_font(r, 22, bold=True, color=NAVY)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(6)
r = p.add_run('Plateforme ArEM — Archives de l\'ENS de Maroua')
set_font(r, 13, italic=True, color=BLUE)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — OUTILS ET TECHNOLOGIES UTILISÉS
# ══════════════════════════════════════════════════════════════════════════════
heading1('1.  Outils et Technologies Utilisés')

body(
    'La plateforme ArEM a été développée en s\'appuyant sur un ensemble cohérent '
    'd\'outils, de frameworks, de langages et de bibliothèques. Cette section présente '
    'chaque composant technologique retenu et justifie son rôle dans l\'architecture '
    'globale de l\'application.'
)

# ── 1.1 Langage de programmation ─────────────────────────────────────────────
heading2('1.1  Langage de Programmation')

heading3('PHP 8.2+')
body(
    'PHP est le langage côté serveur choisi pour le développement du back-end. '
    'La version 8.2 introduit des améliorations significatives en matière de '
    'performances (JIT compiler), de typage strict (readonly properties, intersection '
    'types) et de gestion des erreurs, ce qui renforce la robustesse du code produit.'
)

# ── 1.2 Frameworks ──────────────────────────────────────────────────────────
heading2('1.2  Frameworks')

heading3('Laravel 12.x')
body(
    'Laravel est le framework PHP MVC (Modèle-Vue-Contrôleur) qui constitue le '
    'socle principal d\'ArEM. Il fournit :'
)
for item in [
    'Un système de routage expressif et flexible (routes/web.php).',
    'Un ORM puissant — Eloquent — pour l\'interaction avec la base de données.',
    'Un moteur de templates — Blade — pour la génération des vues HTML.',
    'Un système de migrations et de seeders pour la gestion du schéma de base de données.',
    'Un pipeline de middlewares pour la gestion de l\'authentification et des autorisations.',
    'Des outils intégrés de validation des formulaires, de gestion des fichiers et des notifications.',
    'Une architecture orientée services facilitant la séparation des responsabilités.',
]:
    bullet(item)

heading3('Laravel Breeze 2.3')
body(
    'Laravel Breeze est le kit de démarrage officiel pour l\'authentification. '
    'Il fournit une implémentation légère et personnalisable des fonctionnalités '
    'd\'inscription, de connexion, de réinitialisation de mot de passe et de '
    'vérification d\'adresse e-mail.'
)

heading3('Alpine.js 3.4.2')
body(
    'Alpine.js est un micro-framework JavaScript déclaratif inspiré de Vue.js. '
    'Il est utilisé côté client pour gérer les interactions légères sans avoir '
    'recours à un framework lourd. Il pilote notamment les formulaires multi-étapes, '
    'les menus déroulants et les comportements dynamiques des composants d\'interface.'
)

# ── 1.3 Langages front-end ──────────────────────────────────────────────────
heading2('1.3  Langages et Moteur de Templates Front-end')

heading3('HTML5 / Blade')
body(
    'Blade est le moteur de templates natif de Laravel. Il permet d\'intégrer '
    'dynamiquement des données PHP dans des vues HTML structurées grâce à une '
    'syntaxe concise (@foreach, @if, @include, @component…). L\'ensemble des '
    'quelque 51 templates de l\'application est écrit en Blade.'
)

heading3('CSS3 — Bootstrap 5.3.0 et Bootstrap Icons 1.11.0')
body(
    'Bootstrap 5 est le framework CSS utilisé pour la mise en page responsive, '
    'le système de grille, les composants prêts à l\'emploi (cartes, badges, '
    'boutons, modales, etc.) et la compatibilité multi-navigateurs. '
    'Bootstrap Icons complète la bibliothèque avec plus de 1 800 icônes vectorielles.'
)

heading3('Tailwind CSS 3.1.0')
body(
    'Tailwind CSS est un framework utilitaire qui complète Bootstrap pour affiner '
    'le style des composants avec des classes précises (couleurs, espacements, '
    'typographie). Il est compilé via PostCSS et intégré au pipeline Vite.'
)

heading3('JavaScript (ES Modules)')
body(
    'Le code JavaScript client est structuré en modules ES natifs. '
    'Le fichier app.js initialise Alpine.js et le fichier bootstrap.js configure '
    'Axios pour les requêtes AJAX, notamment la gestion du token CSRF.'
)

# ── 1.4 Outils de build ──────────────────────────────────────────────────────
heading2('1.4  Outils de Build et d\'Outillage')

heading3('Vite 7.0.7')
body(
    'Vite est l\'outil de build front-end retenu pour ArEM. Il assure la compilation '
    'et le bundling des ressources JavaScript et CSS, le rechargement à chaud (HMR) '
    'en développement, et la génération d\'assets optimisés pour la production.'
)

heading3('PostCSS 8.4.31 / Autoprefixer 10.4.2')
body(
    'PostCSS transforme le CSS source (notamment les directives Tailwind) en CSS '
    'standard. Autoprefixer ajoute automatiquement les préfixes propriétaires '
    'nécessaires à la compatibilité navigateur.'
)

heading3('Concurrently 9.0.1')
body(
    'Ce package npm permet de lancer simultanément plusieurs processus en '
    'développement (serveur Vite et Laravel Sail par exemple), simplifiant '
    'ainsi le workflow du développeur.'
)

# ── 1.5 Base de données ──────────────────────────────────────────────────────
heading2('1.5  Systèmes de Gestion de Base de Données')

heading3('SQLite (développement) / MySQL · PostgreSQL (production)')
body(
    'ArEM utilise SQLite en environnement de développement pour sa légèreté '
    '(aucun serveur à installer). En production, la configuration est adaptée '
    'pour MySQL ou PostgreSQL. L\'ORM Eloquent de Laravel abstrait le SGBD '
    'sous-jacent, rendant la migration transparente. La base de données '
    'comprend 9 tables principales gérant les utilisateurs, les documents, '
    'les types de documents, les départements, les métadonnées, les workflows '
    'de validation, les notifications et les statistiques.'
)

# ── 1.6 Bibliothèques PDF ─────────────────────────────────────────────────────
heading2('1.6  Bibliothèques de Traitement PDF')

heading3('barryvdh/laravel-dompdf 3.1')
body(
    'DomPDF est un convertisseur HTML→PDF intégré à Laravel. Il est utilisé '
    'pour générer dynamiquement des documents PDF à partir de templates Blade, '
    'notamment pour l\'export du curriculum vitae (CV) des utilisateurs.'
)

heading3('TCPDF 6.10')
body(
    'TCPDF est une bibliothèque PHP bas niveau permettant la création et la '
    'manipulation avancée de fichiers PDF. Elle est employée comme moteur de '
    'rendu dans certaines opérations de génération de pages de couverture.'
)

heading3('FPDI 2.6 et FPDI-TCPDF Bridge 2.3')
body(
    'FPDI permet d\'importer des pages PDF existantes et de les utiliser comme '
    'gabarit. Associé au bridge TCPDF, il est au cœur du service PdfCoverService '
    'qui génère une page de couverture institutionnelle (logo ENS, identifiant '
    'ArEM, citation recommandée, informations de droits d\'accès) et la '
    'prépose automatiquement au fichier PDF soumis par l\'auteur.'
)

# ── 1.7 Environnement de développement ───────────────────────────────────────
heading2('1.7  Environnement de Développement et Tests')

heading3('Laravel Sail 1.41')
body(
    'Laravel Sail fournit un environnement Docker complet pour le développement '
    'local, incluant PHP, le serveur web et le SGBD. Il garantit la reproductibilité '
    'de l\'environnement entre développeurs.'
)

heading3('Laravel Pint 1.24')
body(
    'Outil de formatage automatique du code PHP selon les conventions PSR-12 '
    'et les standards Laravel, assurant la cohérence du style dans tout le projet.'
)

heading3('Laravel Pail 1.2.2')
body(
    'Outil de surveillance des logs en temps réel en ligne de commande, '
    'facilitant le débogage pendant le développement.'
)

heading3('Laravel Tinker 2.10.1')
body(
    'REPL (Read-Eval-Print Loop) interactif pour Laravel permettant d\'interroger '
    'et de manipuler la base de données et les modèles directement depuis le terminal.'
)

heading3('PHPUnit 11.5.3')
body(
    'Framework de tests unitaires et fonctionnels PHP. ArEM dispose de tests '
    'automatisés couvrant l\'authentification, la gestion des profils, la '
    'réinitialisation de mot de passe et la vérification d\'e-mail.'
)

heading3('FakerPHP/Faker 1.23')
body(
    'Bibliothèque de génération de données factices (noms, e-mails, textes…) '
    'utilisée dans les factories et les seeders pour peupler la base de données '
    'de test.'
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — EXÉCUTION DE L'APPLICATION
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading1('2.  Exécution de l\'Application')

body(
    'Cette section présente l\'application ArEM à travers ses différentes '
    'fonctionnalités, en suivant le parcours naturel d\'un utilisateur depuis '
    'le premier accès jusqu\'aux opérations d\'administration. Chaque '
    'sous-section correspond à un écran ou un groupe d\'écrans cohérents, '
    'et sera illustré par une capture d\'écran.'
)

# ── A. Démarrage ─────────────────────────────────────────────────────────────
heading2('A.  Démarrage de l\'Application')

heading3('A.1  Page d\'Accueil Publique')
body(
    'La page d\'accueil est le premier point de contact de la plateforme. '
    'Elle est accessible sans authentification et présente :'
)
for item in [
    'Un en-tête avec le logo ArEM, le nom de la plateforme et les liens de navigation principaux.',
    'Une bannière de bienvenue avec un champ de recherche rapide.',
    'Des statistiques globales de la plateforme (nombre de documents publiés, d\'auteurs, de départements).',
    'Les documents récemment ajoutés et les documents les plus consultés/téléchargés.',
    'Un pied de page avec les informations institutionnelles (ENS Maroua, MINESUP, Université de Maroua).',
]:
    bullet(item)
placeholder_image('Page d\'accueil publique de la plateforme ArEM')
caption('Figure 1 — Page d\'accueil publique de la plateforme ArEM')

# ── B. Authentification ───────────────────────────────────────────────────────
heading2('B.  Authentification et Gestion des Comptes')

heading3('B.1  Inscription')
body(
    'La page d\'inscription permet à un nouvel utilisateur de créer un compte. '
    'Le formulaire collecte le nom complet, l\'adresse e-mail et le mot de passe '
    '(avec confirmation). Un e-mail de vérification est ensuite envoyé. '
    'À la création, un identifiant unique de la forme AREM-AUTH-2026-XXXXX est '
    'automatiquement attribué.'
)
placeholder_image('Formulaire d\'inscription')
caption('Figure 2 — Formulaire d\'inscription')

heading3('B.2  Connexion')
body(
    'La page de connexion permet à un utilisateur enregistré de s\'authentifier '
    'avec son e-mail et son mot de passe. Une option "Se souvenir de moi" et '
    'un lien vers la réinitialisation du mot de passe sont disponibles.'
)
placeholder_image('Page de connexion')
caption('Figure 3 — Page de connexion')

heading3('B.3  Réinitialisation du Mot de Passe')
body(
    'En cas d\'oubli, l\'utilisateur saisit son adresse e-mail et reçoit un lien '
    'sécurisé permettant de définir un nouveau mot de passe. Le processus respecte '
    'les meilleures pratiques de sécurité (token à usage unique et à durée de vie limitée).'
)
placeholder_image('Formulaire de réinitialisation du mot de passe')
caption('Figure 4 — Formulaire de réinitialisation du mot de passe')

# ── C. Tableau de bord ────────────────────────────────────────────────────────
heading2('C.  Tableau de Bord Utilisateur')

heading3('C.1  Tableau de Bord')
body(
    'Après connexion, l\'utilisateur est redirigé vers son tableau de bord '
    'personnel. Cet espace affiche un récapitulatif de ses activités : '
    'nombre de documents soumis, statut de validation, notifications non lues '
    'et raccourcis vers les actions principales (soumettre un document, '
    'voir mes documents, modifier le profil).'
)
placeholder_image('Tableau de bord de l\'utilisateur connecté')
caption('Figure 5 — Tableau de bord de l\'utilisateur connecté')

# ── D. Gestion du Profil ──────────────────────────────────────────────────────
heading2('D.  Gestion du Profil Utilisateur')

heading3('D.1  Modification du Profil')
body(
    'L\'utilisateur peut mettre à jour ses informations personnelles et '
    'institutionnelles : photo de profil, biographie, département, e-mail '
    'institutionnel, numéro de téléphone, filière, type d\'utilisateur '
    '(Étudiant, Enseignant, Chercheur) et niveau d\'études.'
)
placeholder_image('Page de modification du profil')
caption('Figure 6 — Page de modification du profil')

heading3('D.2  Modification du Mot de Passe')
body(
    'Une section dédiée permet à l\'utilisateur de modifier son mot de passe '
    'en renseignant l\'ancien mot de passe et le nouveau (avec confirmation).'
)
placeholder_image('Formulaire de modification du mot de passe')
caption('Figure 7 — Formulaire de modification du mot de passe')

heading3('D.3  Suppression du Compte')
body(
    'L\'utilisateur peut demander la suppression définitive de son compte. '
    'Une confirmation par mot de passe est requise pour prévenir les '
    'suppressions accidentelles.'
)
placeholder_image('Modal de confirmation de suppression du compte')
caption('Figure 8 — Confirmation de suppression du compte')

# ── E. Curriculum Vitae ───────────────────────────────────────────────────────
heading2('E.  Module Curriculum Vitae (CV)')

heading3('E.1  Rédaction du CV')
body(
    'ArEM intègre un module de gestion de CV académique. L\'utilisateur peut '
    'renseigner :'
)
for item in [
    'Informations personnelles : date et lieu de naissance, nationalité, adresse.',
    'Liens professionnels : LinkedIn, ORCID, Google Scholar.',
    'Parcours académique (formations) — saisie multiple dynamique.',
    'Expériences professionnelles — saisie multiple dynamique.',
    'Compétences, langues parlées, certifications.',
    'Publications et références professionnelles.',
]:
    bullet(item)
placeholder_image('Module de rédaction du CV')
caption('Figure 9 — Module de rédaction du CV')

heading3('E.2  Export du CV en PDF')
body(
    'Le CV complété peut être exporté en fichier PDF via le service DomPDF. '
    'Le document généré présente une mise en page professionnelle aux couleurs '
    'd\'ArEM (bleu marine #0040A0, bleu ciel #5AC8FA).'
)
placeholder_image('Aperçu du CV exporté en PDF')
caption('Figure 10 — Aperçu du CV exporté en PDF')

# ── F. Dépôt de Documents ─────────────────────────────────────────────────────
heading2('F.  Dépôt et Gestion des Documents')

heading3('F.1  Formulaire de Soumission Multi-étapes')
body(
    'La soumission d\'un document se fait via un formulaire guidé en plusieurs '
    'étapes, adapté au type de document choisi :'
)
for item in [
    'Étape 1 — Informations générales : titre, résumé, mots-clés, langue, année.',
    'Étape 2 — Type de document et département : sélection parmi 14 types disponibles.',
    'Étape 3 — Auteurs : ajout de plusieurs auteurs avec leur affiliation.',
    'Étape 4 — Métadonnées spécifiques : directeur de mémoire/thèse, année académique, etc.',
    'Étape 5 — Droits d\'accès : public, restreint ou embargo (avec date).',
    'Étape 6 — Fichier PDF : téléversement du fichier (max 20 Mo).',
]:
    bullet(item)
placeholder_image('Formulaire de soumission — Étape 1 : informations générales')
caption('Figure 11 — Formulaire de soumission (étape 1 : informations générales)')
placeholder_image('Formulaire de soumission — Étape auteurs')
caption('Figure 12 — Formulaire de soumission (étape auteurs)')
placeholder_image('Formulaire de soumission — Étape droits d\'accès et fichier')
caption('Figure 13 — Formulaire de soumission (étape droits d\'accès et téléversement)')

heading3('F.2  Liste de Mes Documents')
body(
    'L\'utilisateur dispose d\'une page listant tous ses dépôts avec leur statut '
    '(Brouillon, En attente, Publié, Rejeté, Révision demandée). '
    'Des actions de modification et de suppression sont disponibles selon le statut.'
)
placeholder_image('Liste des documents soumis par l\'utilisateur')
caption('Figure 14 — Liste des documents de l\'utilisateur avec statuts')

heading3('F.3  Fiche Détaillée d\'un Document')
body(
    'La page de détail d\'un document affiche toutes ses métadonnées : titre, '
    'résumé, auteurs, mots-clés, département, type, année, langue, droits '
    'd\'accès, identifiant ArEM unique (AREM-DOC-ENS-2026-XXXXX), URL '
    'permanente, nombre de consultations et téléchargements.'
)
placeholder_image('Page de détail d\'un document publié')
caption('Figure 15 — Fiche détaillée d\'un document')

heading3('F.4  Téléchargement avec Page de Couverture')
body(
    'Lors du téléchargement, le service PdfCoverService génère automatiquement '
    'une page de couverture institutionnelle et la prépose au PDF original. '
    'Cette couverture inclut le logo ENS, le titre, les auteurs, l\'identifiant '
    'ArEM, l\'URL permanente, la citation recommandée et les informations '
    'de droits d\'accès.'
)
placeholder_image('Page de couverture générée automatiquement pour un document téléchargé')
caption('Figure 16 — Page de couverture institutionnelle générée automatiquement')

# ── G. Parcourir les Documents ────────────────────────────────────────────────
heading2('G.  Navigation et Recherche dans le Catalogue')

heading3('G.1  Parcourir les Documents (Browse)')
body(
    'La page de navigation permet de parcourir le catalogue de documents publiés '
    'avec filtrage par type de document, département ou année. '
    'Les résultats sont paginés (20 par page).'
)
placeholder_image('Page de navigation / browsing du catalogue')
caption('Figure 17 — Navigation dans le catalogue de documents')

heading3('G.2  Recherche Simple')
body(
    'La barre de recherche rapide, accessible depuis n\'importe quelle page, '
    'effectue une recherche en texte libre dans les titres, résumés et '
    'mots-clés des documents publiés, avec affichage de suggestions en temps réel.'
)
placeholder_image('Résultats de la recherche simple')
caption('Figure 18 — Résultats d\'une recherche simple')

heading3('G.3  Recherche Avancée')
body(
    'La recherche avancée offre des filtres combinables :'
)
for item in [
    'Nom de l\'auteur',
    'Type de document',
    'Département / discipline',
    'Plage d\'années (de … à …)',
    'Langue du document',
]:
    bullet(item)
placeholder_image('Interface de recherche avancée avec filtres multiples')
caption('Figure 19 — Formulaire de recherche avancée')

# ── H. Workflow de Validation ─────────────────────────────────────────────────
heading2('H.  Workflow de Validation des Documents')

heading3('H.1  File d\'Attente de Validation')
body(
    'Les modérateurs disposent d\'une vue listant tous les documents en attente '
    'de validation, avec le nom du déposant, le type de document, la date de '
    'soumission et un accès direct à la fiche détaillée.'
)
placeholder_image('File d\'attente de validation (vue modérateur)')
caption('Figure 20 — File d\'attente de validation pour le modérateur')

heading3('H.2  Décision de Validation')
body(
    'Sur la fiche d\'un document en attente, le modérateur peut :'
)
for item in [
    'Approuver — le document est publié immédiatement et une notification est envoyée à l\'auteur.',
    'Rejeter — un motif de rejet est saisi et transmis à l\'auteur.',
    'Demander une révision — des modifications précises sont indiquées à l\'auteur.',
]:
    bullet(item)
body(
    'Un historique complet des actions de validation (avec date, auteur de l\'action '
    'et commentaire) est conservé pour chaque document.',
    indent=0
)
placeholder_image('Interface de décision de validation d\'un document')
caption('Figure 21 — Interface de décision de validation (approbation / rejet / révision)')

heading3('H.3  Notification à l\'Auteur')
body(
    'Après chaque action de validation, l\'auteur reçoit une notification '
    'dans son espace personnel l\'informant du résultat et des éventuels '
    'commentaires du modérateur.'
)
placeholder_image('Notification reçue par l\'auteur après validation')
caption('Figure 22 — Notification de décision de validation reçue par l\'auteur')

# ── I. Notifications ──────────────────────────────────────────────────────────
heading2('I.  Système de Notifications')

heading3('I.1  Centre de Notifications')
body(
    'Un indicateur visuel (badge) dans la barre de navigation signale le nombre '
    'de notifications non lues. Le centre de notifications liste toutes les '
    'alertes avec leur type, leur message et leur date, et permet de les '
    'marquer comme lues individuellement ou globalement.'
)
placeholder_image('Centre de notifications de l\'utilisateur')
caption('Figure 23 — Centre de notifications')

# ── J. Administration ─────────────────────────────────────────────────────────
heading2('J.  Panneau d\'Administration')

heading3('J.1  Tableau de Bord Administrateur')
body(
    'Le tableau de bord administrateur offre une vue synthétique de la '
    'plateforme : nombre total d\'utilisateurs, de documents publiés, de '
    'documents en attente et de départements actifs.'
)
placeholder_image('Tableau de bord de l\'administrateur')
caption('Figure 24 — Tableau de bord administrateur')

heading3('J.2  Gestion des Utilisateurs')
body(
    'L\'administrateur peut consulter la liste complète des utilisateurs, '
    'modifier leurs informations (nom, e-mail, rôle, département) et '
    'supprimer des comptes. L\'affectation des rôles '
    '(Admin, Modérateur, Déposant, Lecteur) se fait directement depuis '
    'cette interface.'
)
placeholder_image('Liste et gestion des utilisateurs')
caption('Figure 25 — Gestion des utilisateurs (liste et modification)')

heading3('J.3  Gestion des Départements')
body(
    'L\'administrateur peut créer, modifier et désactiver des départements. '
    'Chaque département dispose d\'un nom, d\'un code court et d\'une description.'
)
placeholder_image('Gestion des départements')
caption('Figure 26 — Gestion des départements')

heading3('J.4  Gestion des Types de Documents')
body(
    'Cette section permet de configurer les 14 types de documents disponibles '
    'sur la plateforme. Pour chaque type, l\'administrateur définit les champs '
    'obligatoires (liste JSON), le nom, le code et l\'état actif/inactif.'
)
placeholder_image('Gestion des types de documents')
caption('Figure 27 — Gestion des types de documents et de leurs champs requis')

heading3('J.5  Gestion des Documents')
body(
    'La liste de tous les documents déposés sur la plateforme est accessible '
    'à l\'administrateur avec la possibilité de masquer/afficher '
    'un document (toggle is_hidden) sans le supprimer.'
)
placeholder_image('Liste globale des documents (vue administrateur)')
caption('Figure 28 — Gestion globale des documents')

heading3('J.6  Statistiques de la Plateforme')
body(
    'La page de statistiques présente des indicateurs agrégés : '
    'documents par type, par département, évolution des dépôts dans le temps, '
    'documents les plus consultés et les plus téléchargés, avec granularité '
    'journalière pour les vues et téléchargements.'
)
placeholder_image('Tableau de bord des statistiques de la plateforme')
caption('Figure 29 — Statistiques globales de la plateforme')

# ── K. Pages publiques ────────────────────────────────────────────────────────
heading2('K.  Pages Publiques d\'Information')

heading3('K.1  À Propos')
body(
    'La page « À propos » présente la mission d\'ArEM, son contexte institutionnel '
    '(ENS de Maroua), et les objectifs de la plateforme en tant que dépôt '
    'institutionnel d\'archives académiques.'
)
placeholder_image('Page À propos')
caption('Figure 30 — Page À propos de la plateforme')

heading3('K.2  Aide')
body(
    'La page d\'aide guide les utilisateurs à travers les principales '
    'fonctionnalités de la plateforme (comment soumettre un document, '
    'comprendre les statuts, accéder aux ressources, etc.).'
)
placeholder_image('Page d\'aide')
caption('Figure 31 — Page d\'aide')

heading3('K.3  Contact')
body(
    'Le formulaire de contact permet à tout visiteur d\'envoyer un message '
    'à l\'équipe d\'ArEM. Les informations de contact institutionnelles '
    '(adresse e-mail, site web) y sont également affichées.'
)
placeholder_image('Page de contact')
caption('Figure 32 — Page de contact')

# ── Save ───────────────────────────────────────────────────────────────────────
output_path = '/home/user/ArEM/Implementation_et_Resultats.docx'
doc.save(output_path)
print(f'Document saved: {output_path}')
